# This function is not intended to be invoked directly. Instead it will be
# triggered by an orchestrator function.
# Before running this sample, please:
# - create a Durable orchestration function
# - create a Durable HTTP starter function
# - add azure-functions-durable to requirements.txt
# - run pip install -r requirements.txt
from azure.cosmos import CosmosClient
import logging
import json
import time
import re
import io
import datetime
import uuid
import re
import os
import pinecone
import numpy as np
from tqdm.auto import tqdm
import pandas as pd
import openai
from azure.storage.blob import BlobServiceClient
from PyPDF2 import PdfWriter, PdfReader
from azure.core.credentials import AzureKeyCredential
from azure.core._match_conditions import MatchConditions
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.cosmos.exceptions import CosmosAccessConditionFailedError

from typing import List, Tuple, Dict, Any, Optional
from docx import Document

class NotEnoughCreditsError(Exception):
    """Raised when the user does not have enough credits to process the pdf"""




DOLLAR_TO_CREDIT = float(os.getenv("ENV_DOLLAR_TO_CREDIT"))
PRICE_PER_1000_PAGES = float(os.getenv("ENV_PRICE_PER_1000_PAGES"))

# ---------- OCR ----------
endpoint = "https://jorgen-receipt-recognizer.cognitiveservices.azure.com/"
key = os.getenv("ENV_OCR_KEY")


def get_paragraphs(page):
    if page["page_content"] is None:
        return []
    paragraphs = page["page_content"].split("\n\n")
    # remove empty paragraphs
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]

    paragraphs_meta = [
        {
            "page_number": page["page_number"],
            "file_name": page["file_name"],
            "content": paragraph,
        }
        for paragraph in paragraphs
    ]
    return paragraphs_meta

def analyze_read(blob, blob_name, PRICE_PER_1000_PAGES, user_credits):
    reader = PdfReader(io.BytesIO(blob))
    num_pages = len(reader.pages)
    logging.info("Number of pages: " + str(num_pages))
    logging.info("Dollars to pay per 1000 pages: " + str(PRICE_PER_1000_PAGES))
    credits_to_pay = PRICE_PER_1000_PAGES * num_pages / 1000* DOLLAR_TO_CREDIT
    logging.info("Credits to pay: " + str(credits_to_pay) + " credits")
    logging.info("User credits: " + str(user_credits) + " credits")
    if user_credits < credits_to_pay:
        raise NotEnoughCreditsError("Not enough credits to process this pdf: price is " + str(credits_to_pay) + " credits and you have " + str(user_credits) + " credits")

    # pages = [
    #             {
    #                 "page_content": page.extract_text().encode("ascii", "ignore").decode("ascii"),
    #                 "file_name": blob_name,
    #                 "page_number": i,
    #             }
    #             for i, page in enumerate(reader.pages)
    #         ]

    # logging.info("Test page 10: " + pages[10]["page_content"])
    # paragraphs = []
    # for page in pages:
    #     paragraphs.extend(get_paragraphs(page))

    # # Check total length of all paragraphs.
    # # If it is less than 100 chars, assume images and go on
    # total_length = sum([len(paragraph["content"]) for paragraph in paragraphs])
    # if total_length >= 100:
    #     return paragraphs, 0, credits_to_pay, num_pages
    # # else move on
    # logging.info("Total length of all paragraphs is less than 100 chars, assuming images and moving on")
    


    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )
    # price is 1.5 dollars per 1000 pages
    p = 1.5
    logging.info("Cost per 1000 pages: $" + str(p)) 
    all_paragraphs = []
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            poller = document_analysis_client.begin_analyze_document("prebuilt-read", document=blob)
            result = poller.result()
            # If the code got this far without throwing an exception, the operation must have been successful
            break
        except Exception as e:
            logging.error(f"Error in form recognizer endpoint. Attempt {attempt+1} of {max_retries} failed with error: {e}")
            if attempt+1 == max_retries:
                raise  # re-raise the last exception if this was the last attempt
            time.sleep(5*(attempt+1))  # sleep for 5 seconds multiplied by the attempt number


    logging.info("Document contains {} pages: ".format(len(result.pages)))

    cost = p*len(result.pages)/1000

    # logging.info("----Languages detected in the document----")
    # for language in result.languages:
    #     logging.info("Language code: '{}' with confidence {}".format(language.locale, language.confidence))

    for paragraph in result.paragraphs:
        # logging.info(
        #     "...Paragraph of length'{}'".format(
        #         len(paragraph.content)
        #     )
        # )
        if len(paragraph.bounding_regions) > 1:
            # throw exception
            logging.info("Error: more than one bounding region")


        all_paragraphs.append({
            "content": paragraph.content,
            "page_number": paragraph.bounding_regions[0].page_number,# + page_number_base,
            "file_name": blob_name,
            "bounding_box": [[{"x": point.x, "y":point.y} for point in paragraph.bounding_regions[0].polygon]]
        })

    return all_paragraphs, cost, credits_to_pay, num_pages


def check_reference_or_url(paragraph):
    # Add your regular expressions here
    regexes = [
        # APA (journal articles)
        r'([A-Za-z]+, \w\. (, & [A-Za-z]+, \w\. )* \(\d{4}\)\. .+?\. [A-Za-z]+, \d{1,3}\(\d{1,3}\), \d{1,5}\-\d{1,5}\.)',
        # MLA (journal articles)
        r'([A-Za-z]+, [A-Za-z]+ (, and [A-Za-z]+ [A-Za-z]+ )*\. “.+?” .+?, vol\. \d+, no\. \d+, \d{4}, pp\. \d+-\d+.)',
        # Chicago (journal articles)
        r'([A-Za-z]+, [A-Za-z]+ (, and [A-Za-z]+ [A-Za-z]+ )*\. “.+?” .+? \d+, no\. \d+ \(\d{4}\): \d+-\d+.)',
        # Harvard (journal articles)
        r'([A-Za-z]+, [A-Za-z]+ (, and [A-Za-z]+ [A-Za-z]+ )*\. \(\d{4}\) .+? .+?, \d{1,3}\(\d{1,3}\), pp\. \d{1,5}-\d{1,5}\.)',
        # IEEE (journal articles)
        r'(\[[0-9]+\] [A-Za-z]+\. [A-Za-z]+ (, "[A-Za-z]+\. [A-Za-z]+", )* ".*", .+, vol\. \d+, no\. \d+, pp\. \d+-\d+, \d{4}\.)',
        # Vancouver (journal articles)
        r'([A-Za-z]+ [A-Za-z]+ (, [A-Za-z]+ [A-Za-z]+ )*\. .+?\. .+?\. \d{4};\d+:\d+-\d+\.)',
        # URLs
        r'(https?://[^\s]+)',
        # APA (books)
        r'([A-Za-z]+, \w\. (, & [A-Za-z]+, \w\. )* \(\d{4}\)\. .+?\. [A-Za-z]+\.)',
        # APA (online articles)
        r'([A-Za-z]+, \w\. (, & [A-Za-z]+, \w\. )* \(\d{4}, \w+ \d+\)\. .+?\. .+?\. https?://[^\s]+)',
        # Last name, Initial
        r'([A-Za-z]+, \w\.)'
    ]

    
    for regex in regexes:
        if re.search(regex, paragraph):
            return True  # If paragraph matches any regex, return True

    return False  # If paragraph doesn't match any regex, return False


# ---------- Combine and clean paragraphs ----------
def combine_and_clean_paragraphs(paragraphs):
    min_paragraph_length = 75
    cleaned_paragraphs = []

    for paragraph in paragraphs:
        #logging.info(i)
        content = paragraph["content"]
        page_number = paragraph["page_number"]
        if len(content) < min_paragraph_length: #this works OK, but it's not perfect
            #remove it
            # logging.info("Removing: '" + content + "' at page " + str(page_number) + " of file " + paragraph["file_name"])
            continue
    
        # check if the paragraph is a reference or a URL
        if check_reference_or_url(content):
            # remove
            logging.info("Removing reference or URL: '" + content + "' at page " + str(page_number) + " of file " + paragraph["file_name"])
            continue

        else:
            cleaned_paragraphs.append(paragraph)
                    
    logging.info("Number of paragraphs after clean: " +str(len(cleaned_paragraphs)))

    # now check if the last paragraph of each page ends with a period
    # if not, append the next paragraph to it
    i = 0
    while i < len(cleaned_paragraphs):
        content = cleaned_paragraphs[i]["content"]
        page_number = cleaned_paragraphs[i]["page_number"]
        file_name = cleaned_paragraphs[i]["file_name"]

        bounding_box = cleaned_paragraphs[i].get("bounding_box")

        if i < len(cleaned_paragraphs) - 1:
            next_file_name = cleaned_paragraphs[i+1]["file_name"]
            if file_name != next_file_name: # only combine paragraphs from the same file
                i+=1
                continue

            next_page_number = cleaned_paragraphs[i+1]["page_number"]
            if page_number != next_page_number and not content.endswith("."):
                next_content = cleaned_paragraphs[i+1]["content"]
                next_bounding_box = cleaned_paragraphs[i+1].get("bounding_box")
                # append next_content to content
                # logging.info(f"Appending paragraph on page {page_number} of file {file_name} with paragraph on page {next_page_number} of file {next_file_name}")
                cleaned_paragraphs[i]["content"] = content + " " + next_content
                # if bounding_box and next_bounding_box, combine them. else leave it as None
                if bounding_box and next_bounding_box:
                    cleaned_paragraphs[i]["bounding_box"] = bounding_box + next_bounding_box
                else:
                    cleaned_paragraphs[i]["bounding_box"] = None
                cleaned_paragraphs.pop(i+1)
        i += 1

    logging.info("Number of paragraphs after combining paragraphs split across pages: "+ str(len(cleaned_paragraphs)))

    return cleaned_paragraphs



#------------Split long paragraphs into smaller paragraphs----------------
def split_paragraph(paragraph_content, segment_length=500, overlap_length=70):
  # Split the text into a list of sentences
  sentences = re.split(r'[.!?]', paragraph_content)

  # Initialize the list of segments
  segments = []

  # Initialize the current segment with the first sentence
  segment = sentences[0] + '.'

  # Iterate through the rest of the sentences
  for i, sentence in enumerate(sentences[1:]):
    # If adding the sentence to the segment would make it too long,
    # append the current segment to the list of segments and start a new segment
    if len(segment) + len(sentence) > segment_length:
      # Add the overlap by concatenating the first 50 words of the next segment
      overlap = sentence[:overlap_length] + "..."
      segment += overlap
      segments.append(segment)
      # initialize the next segment with words from the previous sentence plus this sentence
      segment = "..." + sentences[i -1][-overlap_length:] + "." + sentence + '.'
    # Otherwise, add the sentence to the current segment
    else:
      segment += sentence + '.'

  # Append the final segment to the list of segments
  segments.append(segment)

  return segments

def split_long_paragraphs(paragraphs):
    new_paragraphs = []
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if len(paragraph["content"]) > 250*6: #250 words ish
            # logging.info(f"...Splitting paragraph with length {len(paragraph['content'])} on page {paragraph['page_number']} of file {paragraph['file_name']}")
            segments = split_paragraph(paragraph["content"], segment_length=min(len(paragraph["content"])/2, 250*5), overlap_length=100)
            # logging.info("Length of new segments: "+ str([len(segment) for segment in segments]))
            # add segments as new paragraphs
            for segment in segments:
                new_paragraphs.append({"page_number": paragraph["page_number"], "content": segment, "bounding_box": paragraph.get("bounding_box"), "file_name": paragraph["file_name"]})

        else:
            new_paragraphs.append(paragraph)


    return new_paragraphs
    

# ---------- Embed paragraphs ----------
# openai.api_type = "azure"
# openai.api_key = os.getenv("AZURE_OPENAI_API_KEY")
# openai.api_base = os.getenv("AZURE_OPENAI_API_INSTANCE_NAME") 
# openai.api_version = os.getenv("AZURE_OPENAI_API_VERSION")
# engine = os.getenv("AZURE_OPENAI_API_DEPLOYMENT_NAME_EMBEDDINGS")

engine = "text-embedding-ada-002"
openai.api_key = os.getenv("ENV_OPENAI_API_KEY")
pinecone.init(
    api_key=os.getenv("ENV_PINECONE_API_KEY"),
    environment=os.getenv("ENV_PINECONE_ENVIRONMENT")
)

def embed_paragraphs(paragraphs, namespace, index_name):
    if index_name not in pinecone.list_indexes():
        pinecone.create_index(index_name, dimension=1536)
        time.sleep(10)
        logging.info("Index created")
    # connect to index
    index = pinecone.Index(index_name)

    # calculate the total number of characters in the document
    total_chars = sum([len(paragraph["content"]) for paragraph in paragraphs])
    price = 0.0004*total_chars/1000
    logging.info("Approx. cost of embedding document: $"+str(price))

    batch_size = 32  # process everything in batches of 32
    for i in tqdm(range(0, len(paragraphs), batch_size)):
        # set end position of batch
        i_end = min(i+batch_size, len(paragraphs))
        # get batch of lines and IDs
        paragraphs_batch = paragraphs[i:i_end]
        ids_batch = ["" + namespace + "_" + str(n) for n in range(i, i_end)] # TODO: make sure that this works even when uploading new ones to the same namespace
        #logging.info("ids_batch", ids_batch)
        # prep metadata and upsert batch
        #for paragraph in paragraphs:
            #logging.info(paragraph["content"][:6])

        meta = []
        for paragraph in paragraphs_batch:
            temp = {
                "page_number": paragraph["page_number"],
                "file_name": paragraph["file_name"],
                "content": paragraph["content"]
            }

            if "---split---" in paragraph["file_name"]:
                split_file_name, page_numbers = paragraph["file_name"].split("---split---")
                start_page, end_page = page_numbers.replace(".pdf", "").split("-")
                temp["file_name"] = split_file_name + ".pdf"
                temp["page_number"] += int(start_page)  # Adding the starting page number to the current page number

            if paragraph["bounding_box"]:
                temp["bounding_box"] = json.dumps(paragraph["bounding_box"])

            meta.append(temp)


        #logging.info(meta)

        content_batch = [paragraph["content"] for paragraph in paragraphs_batch]

        # create embeddings for batch
        retry = True
        while retry:
            try:
                res = openai.Embedding.create(input=content_batch, engine=engine)
                retry = False
            except Exception as e:
                logging.info("Error in embedding: "+ str(e))
                logging.info("Retrying...")
                time.sleep(10)
                retry = True

        embeds = [record['embedding'] for record in res['data']]
        #logging.info(lines_batch["ada_search"])
        to_upsert = list(zip(ids_batch, embeds, meta))

        # get average length of meta["content"]
        #logging.info("Average length of content:", np.mean([len(up[2]["content"]) for up in to_upsert]))
        #logging.info(to_upsert.__class__.__name__)
        #logging.info(type(to_upsert))
        #logging.info(to_upsert)
        # upsert to Pinecone
        index.upsert(vectors=to_upsert, namespace=namespace)    

    return price


#----------- COsmos DB ----------

cosmos_endpoint = "https://nlpcosmos.documents.azure.com:443/"
cosmos_key = os.getenv("ENV_COSMOS_KEY")
cosmos_client = CosmosClient(url=cosmos_endpoint, credential=cosmos_key)
logging.info("Cosmos DB client created")
cosmos_database = cosmos_client.get_database_client("nlp-search") 
cosmos_container = cosmos_database.get_container_client("users")

# Retrieve the connection string for use with the application. The blob storage
connect_str = os.getenv("ENV_AZURE_STORAGE_CONNECTION_STRING")


def extract_text_from_doc(blob, price_per_1000_pages, user_credits) -> Tuple[List[str], int, int, int]:
    """
    Extracts text from a docx file.

    Args:
        blob: The blob object representing the docx file.
        price_per_1000_pages: The price per 1000 pages.
        user_credits: The user's credits.

    Returns:
        A tuple containing a list of all paragraphs in the docx file, the price, credits to pay, and number of pages.
    """
    document = Document(blob)
    paragraphs = []
    for para in document.paragraphs:
        paragraphs.append(para.text)
    num_pages = len(paragraphs)
    price = (num_pages / 1000) * price_per_1000_pages
    credits_to_pay = price * DOLLAR_TO_CREDIT
    if user_credits < credits_to_pay:
        raise NotEnoughCreditsError("User does not have enough credits to process this file.")
    return paragraphs, price, credits_to_pay, num_pages


def extract_text_from_txt(blob, price_per_1000_pages, user_credits) -> Tuple[List[str], int, int, int]:
    """
    Extracts text from a txt file.

    Args:
        blob: The blob object representing the txt file.
        price_per_1000_pages: The price per 1000 pages.
        user_credits: The user's credits.

    Returns:
        A tuple containing a list of all paragraphs in the txt file, the price, credits to pay, and number of pages.
    """
    text = blob.decode('utf-8')
    paragraphs = text.split('\n\n')
    num_pages = len(paragraphs)
    price = (num_pages / 1000) * price_per_1000_pages
    credits_to_pay = price * DOLLAR_TO_CREDIT
    if user_credits < credits_to_pay:
        raise NotEnoughCreditsError("User does not have enough credits to process this file.")
    return paragraphs, price, credits_to_pay, num_pages


def update_cosmos_user(user, namespace, cost, credits_to_pay, num_pages, blob_name):
    while True:
        etag = user["_etag"]

        # update projects total cost and files in project in cosmos db
        projects = user["projects"]

        # subtract price from user credits
        user["credits"] -= credits_to_pay

        for project in projects:
            if project["namespace"] == namespace:
                # update cost
                if "cost" in project:
                    project["cost"] += cost
                else:
                    project["cost"] = cost

                # save info of file

                blob_name = blob_name.split("---split---")[0] + ".pdf" if "---split---" in blob_name else blob_name

                file_info = {
                    "blob_name": blob_name,
                    "price": cost,
                    "credits": credits_to_pay,
                    "num_pages": num_pages,
                    "file_name": blob_name.split("/")[-1],
                }

                if "files" in project:
                    if file_info["blob_name"] not in [file["blob_name"] for file in project["files"]]:
                        project["files"].append(file_info)
                    else:
                        # add price, credits, num_pages to existing file
                        for file in project["files"]:
                            if file["blob_name"] == file_info["blob_name"]:
                                file["price"] += cost
                                file["credits"] += credits_to_pay
                                file["num_pages"] += num_pages
                            
                else:
                    project["files"] = [file_info]

        try:
            cosmos_container.upsert_item(user, etag=etag, match_condition=MatchConditions.IfNotModified)
            break
        except CosmosAccessConditionFailedError:
            user = cosmos_container.read_item(item=user["id"], partition_key=user["id"])
            logging.info("Cosmos DB user modified, retrying...")


def main(settings) -> str:
    retry = True
    while retry:
        try:
            res = openai.Embedding.create(input=["test"], engine=engine)
            retry = False
        except Exception as e:
            logging.info("OPENAI NOT WORKING: "+ str(e))
            logging.info("WAITING TO START PROCESSING... retrying in 10 seconds...")
            time.sleep(10)
            retry = True


    # Create the BlobServiceClient object
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)

    container_name = "users"
    container_client = blob_service_client.get_container_client(container_name)

    # get blob name
    blob_name = settings["blob_name"]
    # get namespace
    namespace = settings["namespace"]
    # get index name
    index_name = settings["index_name"]
    #get user id
    user_id = settings["user_id"]
    # get user credits
    user = cosmos_container.read_item(item=user_id, partition_key=user_id)
    user_credits = user["credits"]

    logging.info("\nDownloading blob " + blob_name)
    blob = container_client.download_blob(blob_name).readall()

    paragraphs, cost, credits_to_pay, num_pages = analyze_read(blob, blob_name, PRICE_PER_1000_PAGES, user_credits)

    # Write paragraphs to a JSON file in the blob storage with the same name and then .json
    try:
        json_blob_name = blob_name.split('.pdf')[0] + '-ocr.json'
        json_blob_client = container_client.get_blob_client(json_blob_name)
        json_blob_client.upload_blob(json.dumps(paragraphs), overwrite=True)
        logging.info("Wrote paragraphs to JSON file at " + json_blob_name)
    except Exception as e:
        logging.warning("Could not write paragraphs to JSON file: " + str(e))

    logging.info(f"Number of paragraphs: {len(paragraphs)}")

    cleaned_paragraphs = combine_and_clean_paragraphs(paragraphs) 

    logging.info(f"Cleaned paragraphs. Number of paragraphs now: {len(cleaned_paragraphs)}")

    split_paragraphs = split_long_paragraphs(cleaned_paragraphs)

    logging.info(f"Split paragraphs. Number of paragraphs now: {len(split_paragraphs)}")

    logging.info(f'Extracted {len(split_paragraphs)} paragraphs')
    # logging.info(f'First paragraph: {split_paragraphs[0]}')
    # logging.info(f'Last paragraph: {split_paragraphs[-1]}')

    logging.info("Now embedding paragraphs")
    embed_price = embed_paragraphs(split_paragraphs, namespace, index_name)
    logging.info("Done embedding paragraphs")

    total_cost = cost + embed_price

    logging.info(f"Total cost for blob {blob_name}: ${total_cost}")

    # update projects total cost and files in prpoject in cosmos db
    update_cosmos_user(user, namespace, total_cost, credits_to_pay, num_pages, blob_name)

    logging.info(f"Deleting blob {blob_name}")
    container_client.delete_blob(blob_name)

    return cost

